"""导出服务 — Markdown → PDF / DOCX"""

import logging
from pathlib import Path
from app.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


def export_pdf(content: str, filename: str) -> str:
    """Markdown → PDF 导出

    Returns:
        导出文件的路径
    """
    export_dir = Path(settings.UPLOAD_DIR) / "exports"
    export_dir.mkdir(parents=True, exist_ok=True)
    filepath = export_dir / f"{filename}.pdf"

    try:
        # Markdown → HTML
        html = _md_to_html(content)

        # HTML → PDF (WeasyPrint)
        from weasyprint import HTML
        HTML(string=html).write_pdf(str(filepath))

        logger.info(f"PDF exported: {filepath}")

    except ImportError:
        logger.warning("WeasyPrint not installed — falling back to HTML")
        filepath = export_dir / f"{filename}.html"
        filepath.write_text(html, encoding="utf-8")
    except Exception as e:
        logger.error(f"PDF export failed: {e}")
        raise

    return str(filepath)


def export_docx(content: str, filename: str) -> str:
    """Markdown → DOCX 导出

    Returns:
        导出文件的路径
    """
    export_dir = Path(settings.UPLOAD_DIR) / "exports"
    export_dir.mkdir(parents=True, exist_ok=True)
    filepath = export_dir / f"{filename}.docx"

    try:
        from docx import Document
        from docx.shared import Inches, Pt

        doc = Document()
        # 设置默认字体
        style = doc.styles["Normal"]
        style.font.size = Pt(11)

        # 简单解析 Markdown 并写入
        lines = content.split("\n")
        for line in lines:
            if line.startswith("# ") and not line.startswith("## "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("| "):
                # 简单表格处理（跳过，在 python-docx 中复杂）
                continue
            elif line.strip() == "":
                continue
            elif line.startswith(">"):
                doc.add_paragraph(line[2:], style="Quote")
            elif line.startswith("---"):
                continue
            elif line.startswith("*") and line.endswith("*") and "生成时间" in line:
                continue
            else:
                doc.add_paragraph(line)

        doc.save(str(filepath))
        logger.info(f"DOCX exported: {filepath}")

    except Exception as e:
        logger.error(f"DOCX export failed: {e}")
        raise

    return str(filepath)


def export_markdown(content: str, filename: str) -> str:
    """导出纯 Markdown 文件"""
    export_dir = Path(settings.UPLOAD_DIR) / "exports"
    export_dir.mkdir(parents=True, exist_ok=True)
    filepath = export_dir / f"{filename}.md"
    filepath.write_text(content, encoding="utf-8")
    return str(filepath)


def _md_to_html(md_content: str) -> str:
    """Markdown → HTML 转换（用于 PDF 导出）"""
    try:
        import markdown
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<style>
  body {{ font-family: 'Noto Sans SC', 'Microsoft YaHei', sans-serif; max-width: 800px; margin: 0 auto; padding: 40px; font-size: 13px; line-height: 1.8; color: #333; }}
  h1 {{ border-bottom: 2px solid #1a73e8; padding-bottom: 8px; }}
  h2 {{ border-bottom: 1px solid #ddd; padding-bottom: 4px; }}
  table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
  th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
  th {{ background: #f5f5f5; }}
  blockquote {{ border-left: 4px solid #1a73e8; padding-left: 16px; color: #666; }}
  code {{ background: #f5f5f5; padding: 2px 6px; border-radius: 3px; }}
  pre {{ background: #f5f5f5; padding: 16px; border-radius: 8px; overflow-x: auto; }}
</style>
</head>
<body>
{markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'codehilite'])}
</body>
</html>"""
    except ImportError:
        return f"<pre>{md_content}</pre>"
